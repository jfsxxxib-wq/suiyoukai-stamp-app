from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "seed-for-95-percent-memory-design-2026-07-05.md"
OUT = ROOT / "dist" / "水曜会_95パーセントの種メモ_2026-07-05.docx"


KEY_LABELS = {
    "理由:",
    "8/30までにやってよいこと:",
    "8/30までに急がないこと:",
    "基本の声かけ:",
    "少し説明する時:",
    "迷った時にスタッフが見る基準:",
    "候補:",
    "使い方の方針:",
    "最小案:",
    "まだ決めないこと:",
    "判断基準:",
    "最初の到達点:",
}


def set_font(run, name="Yu Gothic", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def add_title(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, after=2)
    r = p.add_run(title)
    set_font(r, size=22, bold=True, color="24563A")

    p2 = doc.add_paragraph()
    set_spacing(p2, after=12)
    r2 = p2.add_run("8/30イベント後に、今日の花を思い出として開けるようにするための設計メモ")
    set_font(r2, size=10.5, color="555555")


def add_callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F3F8F0")
    p = cell.paragraphs[0]
    set_spacing(p, after=0)
    r = p.add_run(text)
    set_font(r, size=11, bold=True, color="24563A")
    doc.add_paragraph()


def add_section_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            set_font(run, size=15, bold=True, color="2E6B45")
        elif level == 2:
            set_font(run, size=12.5, bold=True, color="2E6B45")
        else:
            set_font(run, size=11.5, bold=True, color="24563A")
    return heading


def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    set_spacing(p, after=6)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=bold, color="24563A" if bold else None)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4)
    r = p.add_run(text)
    set_font(r, size=10.5)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=4)
    r = p.add_run(text)
    set_font(r, size=10.5)


def configure_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("水曜会 95%の種メモ")
    set_font(r, size=9, color="777777")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_doc(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    title = lines[0].removeprefix("# ").strip()
    add_title(doc, title)

    for raw in lines[1:]:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            add_section_heading(doc, line[3:], 1)
        elif line.startswith("### "):
            add_section_heading(doc, line[4:], 2)
        elif line.startswith("> "):
            add_callout(doc, line[2:])
        elif line.startswith("- "):
            add_bullet(doc, line[2:])
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", "．"):
            add_number(doc, line.split(maxsplit=1)[1] if " " in line else line)
        else:
            add_body(doc, line, bold=line in KEY_LABELS or line.endswith(":"))

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
