from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dist" / "水曜会_8月30日に向けた現在地と作業記録_2026-07-05.docx"
SOURCES = [
    ("8/30に向けた現在地メモ", ROOT / "docs" / "current-status-for-aug30-2026-07-05.md"),
    ("作業記録 2026-07-05", ROOT / "docs" / "work-log-2026-07-05.md"),
]


LABELS = {
    "理由:",
    "今日の合言葉:",
    "当日の声かけ:",
    "判定:",
    "作成・更新した資料:",
    "Word保存:",
    "紙運営で固定したこと:",
    "残した考え方:",
    "守る理由:",
    "守る運用:",
    "今は保留すること:",
    "大きな未解決問題:",
    "小さな注意:",
}


def set_font(run, name="Yu Gothic", size=10.5, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("水曜会 8/30に向けた現在地と作業記録")
    set_font(r, size=9, color="777777")


def title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, after=2)
    r = p.add_run("水曜会 8/30に向けた現在地と作業記録")
    set_font(r, size=21, bold=True, color="24563A")

    p2 = doc.add_paragraph()
    set_spacing(p2, after=12)
    r2 = p2.add_run("2026年7月5日の仕上げ確認と作業ログをまとめた保存用資料")
    set_font(r2, size=10.5, color="555555")


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            set_font(run, size=15, bold=True, color="2E6B45")
        elif level == 2:
            set_font(run, size=12.5, bold=True, color="2E6B45")
        else:
            set_font(run, size=11.5, bold=True, color="24563A")


def callout(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    set_cell_fill(cell, "F3F8F0")
    p = cell.paragraphs[0]
    set_spacing(p, after=0)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=True, color="24563A")
    doc.add_paragraph()


def body(doc, text, bold=False):
    p = doc.add_paragraph()
    set_spacing(p, after=6)
    r = p.add_run(text)
    set_font(r, bold=bold, color="24563A" if bold else None)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=4)
    r = p.add_run(text)
    set_font(r)


def number(doc, text):
    p = doc.add_paragraph(style="List Number")
    set_spacing(p, after=4)
    r = p.add_run(text)
    set_font(r)


def add_markdown(doc, label, path):
    heading(doc, label, 1)
    skip_first = False
    for raw in path.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if not skip_first:
                skip_first = True
                continue
            heading(doc, line[2:], 1)
        elif line.startswith("## "):
            heading(doc, line[3:], 2)
        elif line.startswith("### "):
            heading(doc, line[4:], 3)
        elif line.startswith("> "):
            callout(doc, line[2:])
        elif line.startswith("- "):
            bullet(doc, line[2:])
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", "．"):
            number(doc, line.split(maxsplit=1)[1] if " " in line else line)
        else:
            body(doc, line, bold=line in LABELS or line.endswith(":"))


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure(doc)
    title(doc)

    for index, (label, path) in enumerate(SOURCES):
        if index:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        add_markdown(doc, label, path)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
