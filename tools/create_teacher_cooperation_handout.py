from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output" / "teacher-cooperation-handout"
DOCX_PATH = OUT_DIR / "水曜会_8月5日_先生向け_期間限定2行コメントご協力のお願い_2026-07-27.docx"

FONT = "Yu Gothic"
INK = "24343A"
TEAL = "2B6F6A"
TEAL_DARK = "1E5551"
TEAL_PALE = "EAF5F3"
GOLD_PALE = "FFF7E2"
GRAY = "66757A"
LINE = "B9D4D1"


def set_cellless_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_paragraph_border(paragraph, color=LINE, size="10", space="7"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    for edge_name in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{edge_name}")
        edge.set(qn("w:val"), "single")
        edge.set(qn("w:sz"), size)
        edge.set(qn("w:space"), space)
        edge.set(qn("w:color"), color)
        borders.append(edge)


def set_font(run, size=10.5, bold=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_keep(paragraph, keep_with_next=False, keep_together=True):
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together


def add_text_paragraph(doc, text, size=10.5, bold=False, color=INK,
                       before=0, after=5, line=1.15, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_spacing(p, before, after, line)
    set_keep(p)
    set_font(p.add_run(text), size, bold, color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    set_spacing(p, 6, 3, 1.0)
    set_keep(p, keep_with_next=True)
    set_font(p.add_run(text), 12.5, True, TEAL_DARK)
    return p


def create_bullet_numbering(doc):
    numbering = doc.part.numbering_part.element
    used_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    used_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abs_id = max(used_abs, default=-1) + 1
    num_id = max(used_num, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abs_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "●")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(tabs)
    p_pr.append(ind)
    lvl.append(start)
    lvl.append(num_fmt)
    lvl.append(lvl_text)
    lvl.append(lvl_jc)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abs_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, text, num_id):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    set_spacing(p, 0, 2.5, 1.1)
    set_keep(p)
    set_font(p.add_run(text), 10.2, False, INK)
    return p


def add_example(doc, label, line1, line2):
    p = doc.add_paragraph()
    set_spacing(p, 1, 2, 1.05)
    set_keep(p)
    set_cellless_shading(p, TEAL_PALE)
    set_paragraph_border(p, color=LINE, size="6", space="5")
    r = p.add_run(label + "　")
    set_font(r, 9.5, True, TEAL_DARK)
    r = p.add_run(line1)
    set_font(r, 9.5, False, INK)
    r.add_break()
    set_font(p.add_run(line2), 9.5, False, INK)
    return p


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    heading = doc.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    heading.font.size = Pt(12.5)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(TEAL_DARK)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.line_spacing = 1.0

    # customer_pack-inspired, restrained title stack for a teacher-facing handout.
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(kicker, 0, 1, 1.0)
    set_keep(kicker, keep_with_next=True)
    set_font(kicker.add_run("水曜会からのご案内"), 9.5, True, TEAL)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, 0, 2, 1.0)
    set_keep(title, keep_with_next=True)
    set_font(title.add_run("「一局のご縁帳」"), 20, True, TEAL_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, 0, 8, 1.0)
    set_keep(subtitle, keep_with_next=True)
    set_font(subtitle.add_run("期間限定2行コメント　ご協力のお願い"), 13, True, INK)

    add_text_paragraph(
        doc,
        "「一局のご縁帳」は、水曜会での一局や先生とのご縁を、参加した方が振り返るための記録です。",
        after=4,
    )
    add_text_paragraph(
        doc,
        "今後、小さな試みとして、先生からの短い「2行コメント」を期間限定で掲載することを考えています。"
        "実施に先立ち、内容と安全な進め方をご説明し、後日のご協力をお願いするものです。",
        after=3,
    )

    add_heading(doc, "2行コメントについて")
    num_id = create_bullet_numbering(doc)
    bullets = [
        "コメントを書くかどうかは、先生ご本人の自由です。",
        "書かない場合は、あらかじめ用意した定型文を表示します。",
        "先生に書いていただいた本文を、運営側が代理で編集することはありません。",
        "公開する前に、先生ご本人が内容を確認できます。",
        "公開期間が終わると、通常の文章へ戻ります。",
    ]
    for item in bullets:
        add_bullet(doc, item, num_id)

    add_heading(doc, "定型文の例")
    add_text_paragraph(doc, "コメントを書かない場合は、次のような文章を表示します。", size=9.8, color=GRAY, after=2)
    add_example(
        doc,
        "例1",
        "本日もご参加いただき、ありがとうございました。",
        "これからも一局一局のご縁を大切にしてまいります。",
    )
    add_example(
        doc,
        "例2",
        "皆さまと囲碁を楽しめたことを、うれしく思います。",
        "また水曜会でお会いできることを楽しみにしています。",
    )
    add_example(
        doc,
        "例3",
        "今日の一局が、よい思い出となれば幸いです。",
        "これからも楽しく囲碁を続けていきましょう。",
    )

    notice = add_text_paragraph(
        doc,
        "今回は計画のご説明と、後日のご協力のお願いだけです。まだコメントの入力・通信・公開は行いません。"
        "安全な試験準備が整いましたら、あらためて日時と方法をご説明します。",
        size=9.8,
        bold=False,
        before=4,
        after=5,
        line=1.1,
    )
    set_cellless_shading(notice, GOLD_PALE)
    set_paragraph_border(notice, color="E7C96F", size="8", space="6")

    ask = doc.add_paragraph()
    ask.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(ask, 2, 5, 1.15)
    set_keep(ask)
    set_cellless_shading(ask, TEAL_PALE)
    set_paragraph_border(ask, color=TEAL, size="12", space="8")
    set_font(ask.add_run("今回お伺いしたいこと\n"), 10.5, True, TEAL_DARK)
    set_font(
        ask.add_run("このような小さな試験に、後日ご協力いただけますか。"),
        12,
        True,
        INK,
    )

    add_text_paragraph(
        doc,
        "ご不明な点や心配なことがありましたら、遠慮なくお知らせください。",
        size=9.8,
        color=GRAY,
        after=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(footer_p, 0, 0, 1.0)
    set_font(footer_p.add_run("水曜会"), 8.5, False, GRAY)

    # Prevent Word from adding personal metadata.
    doc.core_properties.title = "「一局のご縁帳」期間限定2行コメント ご協力のお願い"
    doc.core_properties.subject = "先生向け説明資料"
    doc.core_properties.author = "水曜会"
    doc.core_properties.last_modified_by = "水曜会"

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
