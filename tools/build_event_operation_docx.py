from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dist" / "水曜会_8月30日_運営資料_2026-07-05.docx"

SOURCES = [
    ("当日運営1枚メモ", ROOT / "docs" / "event-day-operation-one-page-memo-2026-07-04.md"),
    ("前日準備チェックリスト", ROOT / "docs" / "event-prep-checklist-2026-07-05.md"),
    ("QR印刷確認チェックリスト", ROOT / "docs" / "qr-print-checklist-2026-07-05.md"),
]


def set_font(run, name="Yu Gothic", size=11, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bottom_border(paragraph, color="D9E6DC"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_title(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, after=3)
    run = title.add_run("水曜会 8/30イベント 運営資料")
    set_font(run, size=22, bold=True, color="24563A")
    add_bottom_border(title)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12)
    run = subtitle.add_run("受付QR、先生ごとの固定QR、対局内容QRを迷わず使うための紙運営セット")
    set_font(run, size=11, color="555555")


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), "F3F8F0")
    tc_pr.append(shade)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=2)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color="24563A")
    p2 = cell.add_paragraph()
    set_paragraph_spacing(p2, after=0)
    r2 = p2.add_run(text)
    set_font(r2, size=10.5)
    doc.add_paragraph()


def parse_markdown(doc, label, path):
    lines = path.read_text(encoding="utf-8").splitlines()
    first_heading_skipped = False

    section_heading = doc.add_heading(label, level=1)
    for run in section_heading.runs:
        set_font(run, size=16, bold=True, color="2E6B45")

    for raw in lines:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            if not first_heading_skipped:
                first_heading_skipped = True
                continue
        if line.startswith("## "):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                set_font(run, size=13, bold=True, color="2E6B45")
            continue
        if line.startswith("> "):
            add_callout(doc, "声かけ", line[2:])
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_spacing(p, after=4, line=1.15)
            r = p.add_run(line[2:])
            set_font(r, size=10.5)
            continue
        if len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", "．"):
            p = doc.add_paragraph(style="List Number")
            set_paragraph_spacing(p, after=4, line=1.15)
            r = p.add_run(line.split(maxsplit=1)[1] if " " in line else line)
            set_font(r, size=10.5)
            continue

        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=6, line=1.15)
        r = p.add_run(line)
        if line.endswith(":") or line in ("PC側:", "注意:", "確認すること:", "見せる言葉:", "もう少し説明する時:", "見せる場所:", "先生を間違えた:", "参加スタンプを間違えた:", "不安な時:"):
            set_font(r, size=10.5, bold=True, color="24563A")
        else:
            set_font(r, size=10.5)


def set_document_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)


def add_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("水曜会 8/30イベント 運営資料")
    set_font(r, size=9, color="777777")


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_document_defaults(doc)
    add_title(doc)
    add_callout(
        doc,
        "当日の合言葉",
        "説明より先に、今日の花を見せる。参加者には「今日の花が入りました。」と伝える。",
    )

    for index, (label, path) in enumerate(SOURCES):
        if index:
            doc.add_section(WD_SECTION_START.NEW_PAGE)
        parse_markdown(doc, label, path)

    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
