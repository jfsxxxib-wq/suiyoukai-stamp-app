from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "dist" / "水曜会_花記録_参加者向け案内_ホーム画面から使う_2026-07-05.docx"
QR = ROOT / "tmp" / "app-entry-qr.png"
APP_ICON = ROOT / "assets" / "app-icon-192.png"
APP_URL = "https://jfsxxxib-wq.github.io/suiyoukai-stamp-app/"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color="D8CDAE", size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=130, bottom=100, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_font(run, size=None, bold=None, color=None):
    run.font.name = "Yu Gothic"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=11, bold=False, color="3E3323"):
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return run


def set_para(paragraph, before=0, after=4, line=1.08, align=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_heading(paragraph, text, size=13, color="2F6B4F"):
    set_para(paragraph, before=2, after=5, line=1.05)
    add_text(paragraph, text, size=size, bold=True, color=color)


def add_step_list(cell, steps):
    for idx, step in enumerate(steps, 1):
        p = cell.add_paragraph()
        set_para(p, after=2, line=1.08)
        add_text(p, f"{idx}. ", size=10.4, bold=True, color="2F6B4F")
        add_text(p, step, size=10.4, bold=True if idx == 1 else False, color="3E3323")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.15)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)

    title = doc.add_paragraph()
    set_para(title, after=2, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(title, "水曜会 花記録 参加者案内", size=24, bold=True, color="2F6B4F")

    subtitle = doc.add_paragraph()
    set_para(subtitle, after=7, line=1.05, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(subtitle, "ホーム画面の花アイコンから開いて使います", size=12, bold=True, color="7A5A1B")

    top_table = doc.add_table(rows=1, cols=2)
    top_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(top_table, 10280)
    top_table.columns[0].width = Cm(10.2)
    top_table.columns[1].width = Cm(7.9)
    left, right = top_table.rows[0].cells
    set_cell_width(left, 5780)
    set_cell_width(right, 4500)
    for cell in (left, right):
        set_cell_border(cell, color="D8CDAE", size="8")
        set_cell_margins(cell, top=130, start=150, bottom=130, end=150)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(left, "FFF8E8")
    set_cell_shading(right, "FFFFFF")

    p = left.paragraphs[0]
    set_para(p, after=4, line=1.15)
    add_text(p, "最初にすること", size=14, bold=True, color="2F6B4F")
    p = left.add_paragraph()
    set_para(p, after=4, line=1.25)
    add_text(p, "1. 右のQRコードを読み取ります。\n2. iPhoneはSafari、AndroidはChromeで開きます。\n3. ホーム画面に「水曜会 花記録」を追加します。", size=11.5, bold=True, color="3E3323")
    p = left.add_paragraph()
    set_para(p, before=2, after=0, line=1.2)
    add_text(p, "追加した後は、ホーム画面の花アイコンから開いて使います。", size=10.2, bold=True, color="7A5A1B")

    qr_p = right.paragraphs[0]
    set_para(qr_p, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    if QR.exists():
        qr_p.add_run().add_picture(str(QR), width=Cm(5.5))
    label = right.add_paragraph()
    set_para(label, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(label, "アプリ入口QR", size=12, bold=True, color="2F6B4F")
    url = right.add_paragraph()
    set_para(url, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(url, APP_URL, size=7.8, color="555555")

    doc.add_paragraph()

    steps_table = doc.add_table(rows=1, cols=2)
    steps_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(steps_table, 10280)
    ios_cell, android_cell = steps_table.rows[0].cells
    set_cell_width(ios_cell, 5140)
    set_cell_width(android_cell, 5140)
    for cell in (ios_cell, android_cell):
        set_cell_border(cell, color="D8CDAE", size="8")
        set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
        set_cell_shading(cell, "FFFFFF")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    p = ios_cell.paragraphs[0]
    add_heading(p, "iPhoneの人")
    add_step_list(ios_cell, [
        "QRコードを読み取る",
        "Safariで開く",
        "共有ボタンを押す",
        "ホーム画面に追加",
        "名前が「水曜会 花記録」なら追加",
        "次からは花アイコンで開く",
    ])

    p = android_cell.paragraphs[0]
    add_heading(p, "Androidの人")
    add_step_list(android_cell, [
        "QRコードを読み取る",
        "Chromeで開く",
        "右上のメニューを押す",
        "ホーム画面に追加、またはアプリをインストール",
        "名前が「水曜会 花記録」なら追加",
        "次からは花アイコンで開く",
    ])

    doc.add_paragraph()

    note_table = doc.add_table(rows=1, cols=1)
    note_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(note_table, 10280)
    note = note_table.rows[0].cells[0]
    set_cell_width(note, 10280)
    set_cell_border(note, color="C7D8C6", size="8")
    set_cell_margins(note, top=120, start=170, bottom=120, end=170)
    set_cell_shading(note, "F2F8EF")
    p = note.paragraphs[0]
    set_para(p, after=2, line=1.15)
    add_text(p, "大事: 追加できたら、必ずホーム画面の花アイコンから開いてください。", size=12, bold=True, color="2F6B4F")
    p = note.add_paragraph()
    set_para(p, after=3, line=1.2)
    add_text(p, "Edgeや別のブラウザで開いた画面は、ホーム画面の記録と分かれることがあります。QRを読んで別画面が開いた時は、花アイコンの「水曜会 花記録」に戻ってください。", size=10.2, bold=True, color="6F3F2C")
    p = note.add_paragraph()
    set_para(p, after=3, line=1.2)
    add_text(p, "参加フォームでは、アプリに出ている8桁の受付番号を入力します。「番号をコピーしてフォームを開く」を押すと貼り付けできます。", size=10.2, color="3E3323")
    p = note.add_paragraph()
    set_para(p, after=0, line=1.2)
    add_text(p, "分からない時は、近くのスタッフにこの紙を見せてください。", size=10.5, bold=True, color="3E3323")

    footer = doc.add_paragraph()
    set_para(footer, before=7, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(footer, "水曜会 8/30 イベント用", size=9, color="777777")

    doc.save(OUT)


if __name__ == "__main__":
    build()
